#!/usr/bin/env python3
"""Generate editable Word report tables for Kopiku Mana."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "laporan_kopiku_mana_tabel.docx"


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]]):
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    document.add_paragraph("")
    return table


def currency(value: int) -> str:
    return f"Rp {value:,.0f}".replace(",", ".")


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_heading("Laporan Pendukung Aplikasi Kopiku Mana", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Deskripsi aktor dan fitur, anggaran pengembangan, serta struktur database.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run("Gambaran singkat: ").bold = True
    intro.add_run(
        "Kopiku Mana adalah aplikasi rekomendasi cafe wilayah Tapal Kuda. "
        "Pengguna dapat mencari cafe, melihat detail, menyimpan wishlist, menulis ulasan, "
        "mengumpulkan poin, memakai referral, top up poin, dan menukar poin untuk akses premium."
    )

    actor_rows = [
        [
            "Pengguna",
            "Aktor utama yang memakai aplikasi untuk mencari dan mengelola rekomendasi cafe.",
            "Login, registrasi, cari & filter cafe, lihat cafe terdekat, lihat detail cafe, wishlist, tulis review, referral, top up poin, upgrade premium, lihat riwayat, atur notifikasi.",
        ],
        [
            "Pengguna Baru",
            "Calon pengguna yang belum memiliki akun dan masuk melalui proses registrasi.",
            "Registrasi akun, menggunakan kode referral opsional, login setelah akun dibuat.",
        ],
        [
            "Pengguna Premium",
            "Bukan aktor terpisah, melainkan status akun pengguna setelah menukar poin premium.",
            "Akses Hidden Gems, wishlist unlimited, badge premium, dan benefit premium lain.",
        ],
        [
            "Mitra Cafe",
            "Pemilik atau pengelola cafe yang ingin bekerja sama dengan Kopiku Mana.",
            "Mengajukan sponsor melalui email/fitur partnership; cafe sponsor dapat tampil lebih prioritas.",
        ],
        [
            "Admin Kopiku Mana",
            "Pihak pengelola sistem yang mengatur data operasional aplikasi.",
            "Mengelola data cafe, status sponsor, top pick, hidden gems, dan data pendukung melalui Firestore/operasional luar aplikasi.",
        ],
        [
            "Firebase Auth",
            "Sistem eksternal untuk autentikasi akun.",
            "Memproses login, registrasi, reset password, dan session pengguna.",
        ],
        [
            "Cloud Firestore",
            "Database utama aplikasi berbasis cloud.",
            "Menyimpan data user, cafe, review, wishlist, poin, referral, dan transaksi.",
        ],
        [
            "GPS / Geolocator",
            "Sistem eksternal perangkat untuk membaca lokasi pengguna.",
            "Mengirim koordinat lokasi agar aplikasi dapat menampilkan cafe terdekat.",
        ],
        [
            "ImgBB",
            "Layanan eksternal untuk penyimpanan gambar.",
            "Menyimpan foto profil dan foto review, lalu mengembalikan URL gambar.",
        ],
        [
            "Backend Railway / Midtrans",
            "Backend dan payment gateway untuk transaksi top up.",
            "Membuat transaksi top up poin, order ID, redirect URL pembayaran, dan status transaksi.",
        ],
        [
            "SharedPreferences",
            "Penyimpanan lokal perangkat.",
            "Menyimpan timestamp kunjungan terakhir, reminder review, dan status baca notifikasi lokal.",
        ],
    ]
    add_table(doc, "1. Deskripsi Aktor dan Fitur", ["Aktor/Sistem", "Deskripsi", "Fitur/Interaksi"], actor_rows)

    budget_items = [
        ["1", "Analisis kebutuhan & perancangan sistem", "Analisis fitur, aktor, alur sistem, use case, DFD, dan struktur database.", 1500000],
        ["2", "Desain UI/UX aplikasi", "Desain tampilan splash, login, home, explore, detail cafe, wishlist, profil, premium, referral, dan top up.", 2500000],
        ["3", "Pengembangan aplikasi Flutter", "Implementasi screen utama, navigasi, komponen UI, filter, wishlist, review, premium, referral, dan profil.", 7500000],
        ["4", "Integrasi Firebase Auth & Firestore", "Login/register, reset password, user profile, cafe, review, wishlist, point history, dan referral history.", 3500000],
        ["5", "Integrasi lokasi GPS/Geolocator", "Permission lokasi, pengambilan koordinat, dan perhitungan jarak cafe terdekat.", 1500000],
        ["6", "Integrasi upload gambar ImgBB", "Upload foto profil dan foto review, penyimpanan URL gambar.", 1250000],
        ["7", "Integrasi top up Railway/Midtrans", "Pembuatan transaksi top up, pengiriman data paket, order ID, dan redirect pembayaran.", 2500000],
        ["8", "Fitur poin, premium, referral, dan sponsor", "Logika poin review/referral, redeem premium, benefit premium, dan prioritas cafe sponsor.", 2500000],
        ["9", "Testing & debugging", "Uji login, pencarian, detail cafe, wishlist, review, premium, referral, top up, dan validasi data.", 2000000],
        ["10", "Dokumentasi laporan & diagram", "Pembuatan laporan, tabel database, use case diagram, DFD level 0, DFD level 1, dan penjelasan.", 1500000],
        ["11", "Deployment & konfigurasi awal", "Konfigurasi Firebase, backend, build aplikasi, dan pengecekan rilis awal.", 1500000],
        ["12", "Cadangan/maintenance awal", "Perbaikan minor, penyesuaian fitur, dan support awal setelah demo/rilis.", 2000000],
    ]
    budget_rows = []
    total = 0
    for no, item, detail, cost in budget_items:
        total += cost
        budget_rows.append([no, item, detail, currency(cost)])
    budget_rows.append(["", "TOTAL ESTIMASI", "Anggaran dapat disesuaikan dengan skala tim, durasi pengerjaan, dan kebutuhan deployment.", currency(total)])
    add_table(doc, "2. Anggaran Pengembangan", ["No", "Kebutuhan", "Rincian", "Estimasi Biaya"], budget_rows)

    note = doc.add_paragraph()
    note.add_run("Catatan anggaran: ").bold = True
    note.add_run(
        "Nominal di atas adalah estimasi untuk kebutuhan laporan/prototype aplikasi. "
        "Biaya dapat dinaikkan atau diturunkan sesuai standar kampus, vendor, jumlah anggota tim, dan cakupan deployment."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)

    db_summary_rows = [
        ["users", "Menyimpan profil pengguna, status premium, total poin, referral, wishlist, dan preferensi.", "UID pengguna"],
        ["cafes", "Menyimpan data cafe, lokasi, rating, kategori, harga, fasilitas, dan status sponsor/kurasi.", "cafeId"],
        ["reviews", "Menyimpan ulasan pengguna terhadap cafe beserta rating dan foto opsional.", "reviewId"],
        ["point_history", "Menyimpan riwayat perubahan poin seperti referral, review, redeem premium, dan top up.", "historyId"],
        ["referral_history", "Menyimpan riwayat penggunaan kode referral antar pengguna.", "referralHistoryId"],
        ["topup_transactions", "Menyimpan transaksi top up poin dari backend pembayaran.", "orderId/transactionId"],
        ["SharedPreferences", "Penyimpanan lokal untuk notifikasi, bukan collection Firestore.", "local key"],
    ]
    add_table(doc, "3.1 Ringkasan Struktur Database", ["Collection/Data Store", "Fungsi", "Primary Key"], db_summary_rows)

    db_detail_rows = [
        ["users", "uid", "String", "ID unik pengguna dari Firebase Auth."],
        ["users", "name", "String", "Nama pengguna."],
        ["users", "email", "String", "Email akun pengguna."],
        ["users", "photoUrl / avatarUrl", "String", "URL foto profil pengguna dari ImgBB/Firebase profile."],
        ["users", "isPremium", "Boolean", "Status apakah pengguna premium."],
        ["users", "premiumExpiry", "Timestamp/String Date", "Tanggal kedaluwarsa premium."],
        ["users", "totalPoints", "Number", "Jumlah poin pengguna."],
        ["users", "referralCode", "String", "Kode referral milik pengguna."],
        ["users", "hasUsedReferral / usedReferralCode / referredBy", "Boolean/String", "Status dan kode referral yang pernah digunakan."],
        ["users", "wishlist", "Array<String>", "Daftar cafeId yang disimpan pengguna."],
        ["users", "createdAt", "Timestamp", "Tanggal akun dibuat."],
        ["cafes", "name", "String", "Nama cafe."],
        ["cafes", "city", "String", "Kota wilayah Tapal Kuda."],
        ["cafes", "district", "String", "Kecamatan/lokasi area cafe."],
        ["cafes", "address", "String", "Alamat cafe."],
        ["cafes", "photos", "Array<String>", "Daftar URL foto cafe."],
        ["cafes", "priceRange", "String", "Kategori harga: Budget, Mid, Premium."],
        ["cafes", "atmosphere", "Array<String>", "Tag suasana seperti cozy, outdoor, nugas, dan lain-lain."],
        ["cafes", "category", "Array<String>", "Kategori cafe/fasilitas/kebutuhan pengguna."],
        ["cafes", "hasWifi", "Boolean", "Status ketersediaan wifi."],
        ["cafes", "rating / averageRating", "Number", "Rata-rata rating cafe."],
        ["cafes", "totalReviews", "Number", "Jumlah review cafe."],
        ["cafes", "status", "String", "Status cafe: aktif, tutup, atau perlu dikonfirmasi."],
        ["cafes", "isTopPick", "Boolean", "Penanda cafe rekomendasi top pick."],
        ["cafes", "isHiddenGem", "Boolean", "Penanda cafe hidden gem."],
        ["cafes", "latitude / longitude", "Number", "Koordinat lokasi cafe."],
        ["cafes", "isSponsored", "Boolean", "Status cafe sponsor."],
        ["cafes", "sponsorUntil", "Timestamp", "Tanggal berakhir sponsor."],
        ["cafes", "sponsorPriority", "Number", "Prioritas tampil cafe sponsor."],
        ["reviews", "cafeId", "String", "ID cafe yang diulas."],
        ["reviews", "userId", "String", "ID pengguna pembuat ulasan."],
        ["reviews", "userName", "String", "Nama pengguna yang tampil pada review."],
        ["reviews", "rating", "Number", "Nilai rating dari pengguna."],
        ["reviews", "text", "String", "Isi ulasan pengguna."],
        ["reviews", "photos", "Array<String>", "URL foto ulasan yang diupload ke ImgBB."],
        ["reviews", "isVerified", "Boolean", "Status review valid/terverifikasi."],
        ["reviews", "pointsAwarded", "Number", "Poin yang diberikan dari review."],
        ["reviews", "createdAt", "Timestamp", "Tanggal review dibuat."],
        ["point_history", "userId", "String", "ID pengguna pemilik transaksi poin."],
        ["point_history", "type", "String", "Jenis transaksi poin: referral, redeem, review, topup."],
        ["point_history", "amount", "Number", "Jumlah poin masuk atau keluar."],
        ["point_history", "description", "String", "Keterangan transaksi poin."],
        ["point_history", "createdAt", "Timestamp", "Tanggal transaksi poin."],
        ["referral_history", "referrerId", "String", "ID pengguna pemberi kode referral."],
        ["referral_history", "refereeId", "String", "ID pengguna yang memakai kode referral."],
        ["referral_history", "code", "String", "Kode referral yang digunakan."],
        ["referral_history", "pointsAwarded", "Number", "Jumlah poin bonus referral."],
        ["referral_history", "createdAt", "Timestamp", "Tanggal referral berhasil."],
        ["topup_transactions", "userId", "String", "ID pengguna yang melakukan top up."],
        ["topup_transactions", "orderId", "String", "ID order pembayaran dari backend/Midtrans."],
        ["topup_transactions", "points / basePoints / bonusPoints", "Number", "Jumlah poin yang dibeli dan bonus."],
        ["topup_transactions", "price", "Number", "Nominal harga top up."],
        ["topup_transactions", "paymentMethod", "String", "Metode pembayaran yang dipilih."],
        ["topup_transactions", "status", "String", "Status transaksi: pending, success, failed, dan sejenisnya."],
        ["SharedPreferences", "last_visit", "Integer timestamp", "Waktu terakhir pengguna membuka aplikasi."],
        ["SharedPreferences", "last_review_reminder", "Integer timestamp", "Waktu reminder review terakhir muncul."],
        ["SharedPreferences", "notif_read", "Boolean", "Status apakah notifikasi sudah dibaca."],
    ]
    add_table(doc, "3.2 Detail Field Database", ["Collection/Data Store", "Field", "Tipe Data", "Keterangan"], db_detail_rows)

    closing = doc.add_paragraph()
    closing.add_run("Kesimpulan: ").bold = True
    closing.add_run(
        "Struktur data Kopiku Mana berpusat pada collection users, cafes, reviews, point_history, referral_history, "
        "dan topup_transactions. Data lokal SharedPreferences digunakan untuk kebutuhan notifikasi sederhana di perangkat."
    )

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
